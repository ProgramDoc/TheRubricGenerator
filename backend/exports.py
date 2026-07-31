"""Export formatters for lab session content.

Supports: Word (.docx), LaTeX (.tex), Excel (.xlsx), CSV, Python (.py), R (.R)
"""

import csv
import io
import re
from typing import Any


def export_docx(content: str, title: str) -> bytes:
    """Generate a .docx file from markdown content."""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    doc.core_properties.title = title

    # Title
    heading = doc.add_heading(title, level=1)
    heading.style.font.size = Pt(18)

    # Convert markdown-ish content to docx paragraphs
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Headers
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            # Strip bold markers for clean text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            doc.add_paragraph(clean)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_latex(content: str, title: str) -> str:
    """Convert markdown content to a basic LaTeX document."""
    # Escape LaTeX special chars
    def esc(s):
        for ch in ["&", "%", "$", "#", "_", "{", "}"]:
            s = s.replace(ch, "\\" + ch)
        return s

    lines = [
        r"\documentclass[12pt]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{geometry}",
        r"\geometry{margin=1in}",
        r"\usepackage{hyperref}",
        "",
        r"\title{" + esc(title) + "}",
        r"\date{\today}",
        "",
        r"\begin{document}",
        r"\maketitle",
        "",
    ]

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            lines.append("")
            continue

        if line.startswith("### "):
            lines.append(r"\subsubsection{" + esc(line[4:]) + "}")
        elif line.startswith("## "):
            lines.append(r"\subsection{" + esc(line[3:]) + "}")
        elif line.startswith("# "):
            lines.append(r"\section{" + esc(line[2:]) + "}")
        elif line.startswith("- ") or line.startswith("* "):
            lines.append(r"\begin{itemize}")
            lines.append(r"\item " + esc(line[2:]))
            lines.append(r"\end{itemize}")
        else:
            # Strip markdown formatting
            clean = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", line)
            clean = re.sub(r"\*(.+?)\*", r"\\textit{\1}", clean)
            clean = re.sub(r"`(.+?)`", r"\\texttt{\1}", clean)
            lines.append(esc(clean) if clean == line else clean)

    lines.append("")
    lines.append(r"\end{document}")
    return "\n".join(lines)


def export_xlsx(data: list[dict], title: str) -> bytes:
    """Generate an .xlsx file from a list of dicts."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name limit

    if not data:
        ws.append(["No data available"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    # Headers — union of every row's keys, in first-seen order. Taking them from
    # row 0 alone silently drops any column that first appears in a later row
    # (identical for uniform data, which is the common case).
    headers = list(dict.fromkeys(k for row in data for k in row))
    ws.append(headers)
    header_fill = PatternFill(start_color="274472", end_color="274472", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    for col_idx, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font

    # Data rows
    for row_data in data:
        row = []
        for h in headers:
            val = row_data.get(h, "")
            if isinstance(val, (dict, list)):
                val = str(val)
            row.append(val)
        ws.append(row)

    # Auto-width
    for col_idx, header in enumerate(headers, 1):
        max_len = len(str(header))
        for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, min_row=2):
            for cell in row:
                if cell.value:
                    max_len = max(max_len, min(len(str(cell.value)), 50))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = max_len + 2

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_csv(data: list[dict]) -> str:
    """Generate CSV string from a list of dicts."""
    if not data:
        return "No data available\n"
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=list(data[0].keys()))
    writer.writeheader()
    for row in data:
        # Flatten nested values
        flat = {}
        for k, v in row.items():
            flat[k] = str(v) if isinstance(v, (dict, list)) else v
        writer.writerow(flat)
    return output.getvalue()


def export_python_script(code_blocks: list[dict], *, preamble: list[str] | None = None,
                         title: str = "Statistical Analysis Script") -> str:
    """Generate a Python script from code blocks extracted from chat.

    ``preamble`` overrides the default import lines (used by Synthesis to emit
    a numpy/scipy meta-analysis preamble instead of the Lab's pandas one).
    """
    default_imports = [
        "import pandas as pd",
        "import numpy as np",
        "import matplotlib.pyplot as plt",
        "import seaborn as sns",
        "from scipy import stats",
    ]
    lines = [
        '"""',
        title,
        "Generated by The AI Researcher Lab",
        '"""',
        "",
        *(preamble if preamble is not None else default_imports),
        "",
        "# ─────────────────────────────────────────────",
        "# Analysis Code",
        "# ─────────────────────────────────────────────",
        "",
    ]

    if code_blocks:
        for i, block in enumerate(code_blocks):
            desc = block.get("description", f"Analysis block {i + 1}")
            code = block.get("code", "# No code provided")
            lines.append(f"# {desc}")
            lines.append(code)
            lines.append("")
    else:
        lines.append("# No code blocks found in the conversation.")
        lines.append("# Re-run your analysis with the AI Statistician to generate code.")

    return "\n".join(lines)


def export_r_script(code_blocks: list[dict], *, preamble: list[str] | None = None,
                    title: str = "Statistical Analysis Script") -> str:
    """Generate an R script from code blocks extracted from chat.

    ``preamble`` overrides the default ``library()`` lines (used by Synthesis
    to load ``meta``/``metafor`` instead of the Lab's tidyverse).
    """
    default_libs = [
        "library(tidyverse)",
        "library(ggplot2)",
    ]
    lines = [
        f"# {title}",
        "# Generated by The AI Researcher Lab",
        "",
        *(preamble if preamble is not None else default_libs),
        "",
        "# ─────────────────────────────────────────────",
        "# Analysis Code",
        "# ─────────────────────────────────────────────",
        "",
    ]

    if code_blocks:
        for i, block in enumerate(code_blocks):
            desc = block.get("description", f"Analysis block {i + 1}")
            code = block.get("code", "# No code provided")
            lines.append(f"# {desc}")
            lines.append(code)
            lines.append("")
    else:
        lines.append("# No code blocks found in the conversation.")
        lines.append("# Re-run your analysis with the AI Statistician to generate code.")

    return "\n".join(lines)
